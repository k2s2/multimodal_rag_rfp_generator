import os
from dotenv import load_dotenv
from pathlib import Path
import io
import re
from docx import Document as docx_Doc
import tempfile
import shutil
import streamlit as st

# File handling imports
from file_handler.pdf_handler import extract_pdf_text
from file_handler.docx_handler import extract_docx_text
from file_handler.csv_handler import extract_csv_text
from file_handler.txt_handler import extract_txt_text

# RAG imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.schema import Document
import time

# DB imports
from langchain_community.utilities import SQLDatabase
from langchain_community.agent_toolkits import SQLDatabaseToolkit
import pandas as pd
from sqlalchemy import create_engine, text

# Agent imports
from langchain.agents import AgentExecutor, create_react_agent
from langchain.prompts import PromptTemplate
from langchain.tools import Tool

# LLM imports - using Groq API
from langchain_groq import ChatGroq


load_dotenv()

# Use temp directory for cloud deployment
TEMP_DIR = tempfile.gettempdir()
FOLDER = os.path.join(TEMP_DIR, "rfp_data")
CHROMA_DIR = os.path.join(TEMP_DIR, "chroma_db")

# Ensure directories exist
Path(FOLDER).mkdir(parents=True, exist_ok=True)
Path(CHROMA_DIR).mkdir(parents=True, exist_ok=True)

DB_CONFIG = {
    "host": st.secrets.get("MYSQL_HOST") or os.getenv("MYSQL_HOST"),
    "user": st.secrets.get("MYSQL_USER") or os.getenv("MYSQL_USER"),
    "password": st.secrets.get("MYSQL_PASSWORD") or os.getenv("MYSQL_PASSWORD"),
    "database": st.secrets.get("MYSQL_DATABASE") or os.getenv("MYSQL_DATABASE"),
    "port": int(st.secrets.get("MYSQL_PORT")) or int(os.getenv("MYSQL_PORT"))
}

def create_ollama_llm():
    """Create and return an LLM instance using API"""
    api_key = st.secrets.get("GROQ_API_KEY") or os.getenv("GROQ_API_KEY")
    model_name = "llama-3.3-70b-versatile"
    return ChatGroq(model=model_name, groq_api_key=api_key)

def create_chunks(file):
    file_chunks = []
    try:
        file_extension=file.suffix.lower()
        if file_extension==".pdf":
            full_text=extract_pdf_text(str(file))
        elif file_extension==".docx":
            full_text=extract_docx_text(str(file))
        elif file_extension==".csv":
            full_text=extract_csv_text(str(file))
        elif file_extension==".txt":
            full_text=extract_txt_text(str(file))
        elif file_extension in [".jpg",".jpeg",".png",".gif",".bmp",".tiff"]:
            from file_handler.image_handler import extract_image_text_easyocr,extract_image_description
            full_text=extract_image_text_easyocr(str(file))
            if not full_text or full_text.strip()=="":
                full_text=extract_image_description(str(file))
        elif file_extension in [".mp3",".wav",".flac",".m4a",".ogg"]:
            from file_handler.audio_handler import extract_audio_text
            full_text=extract_audio_text(str(file))
        elif file_extension in [".mp4",".mov",".avi"]:
            from file_handler.video_handler import extract_video_content
            full_text=extract_video_content(str(file))
        else:
            print(f"  Unsupported file type: {file_extension}")
            return file_chunks
        # Creating document
        doc=Document(page_content=full_text, metadata={"source": file.name, "file_path": str(file)})
        text_splitter=RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        # Creating chunks
        chunks=text_splitter.split_documents([doc])
        # Adding metadata to chunks
        for chunk in chunks:
            chunk.metadata["source"] = file.name
            chunk.metadata["file_path"] = str(file)
        file_chunks.extend(chunks)
        print(f"  Created {len(chunks)} chunks from {file.name}")
    except Exception as e:
        print(f"  Error processing {file.name}: {str(e)}")
    return file_chunks

def create_vector_db(all_chunks):
    """Create embeddings and vector store"""
    print("Creating embeddings and vector store...")
    # embeddings = HuggingFaceEmbeddings(model_name="./models/all-MiniLM-L6-v2")
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    # Create a collection
    collection_name = f"file_collection_{int(time.time())}"
    vector_store = Chroma(
        collection_name=collection_name,
        embedding_function=embeddings,
        persist_directory=CHROMA_DIR
    )
    vector_store.add_documents(all_chunks)
    return vector_store

def setup_rag_vectorstore(directory_path):
    """Setup simple RAG with all files in directory and return RAG chain"""
    directory_path = Path(directory_path)
    # Find all files in directory
    file_extensions = ["*.pdf", "*.docx", "*.csv", "*.txt",
                       "*.jpg", "*.jpeg", "*.png", "*.gif", "*.bmp", "*.tiff",
                       "*.mp3", "*.wav", "*.flac", "*.m4a", "*.ogg",
                       "*.mp4", "*.mov", "*.avi"]

    all_files = []
    for extension in file_extensions:
        all_files.extend(directory_path.glob(extension))
    if not all_files:
            raise ValueError(f"No files found in directory: {directory_path}")
    print(f"Found {len(all_files)} files in directory")
    # Process each file
    all_chunks = []
    for file in all_files:
        print(f"Processing: {file.name}")
        # Extract text from current file
        all_chunks.extend(create_chunks(file))
    if not all_chunks:
        raise ValueError("No content could be extracted from any files")
    print(f"Total chunks created: {len(all_chunks)}")
    vectorstore = create_vector_db(all_chunks)
    return vectorstore

def setup_sqlalchemy_engine():
    """Setup SQLAlchemy engine for direct queries"""
    db_uri = f"mysql+pymysql://{DB_CONFIG['user']}:{DB_CONFIG['password']}@{DB_CONFIG['host']}:{DB_CONFIG['port']}/{DB_CONFIG['database']}"
    return create_engine(db_uri)

def execute_direct_query(engine, query):
    """Execute a direct SQL query and return results"""
    with engine.connect() as conn:
        result = conn.execute(text(query))
        if query.strip().upper().startswith(('SELECT', 'SHOW', 'DESCRIBE', 'EXPLAIN')):
            rows = result.fetchall()
            columns = result.keys()
            return pd.DataFrame(rows, columns=columns) if rows else pd.DataFrame()
        else:
            return {"rows_affected": result.rowcount}

def interpret_query_results(llm, original_question, query_result):
    """Use LLM to interpret query results and provide meaningful answer"""
    if isinstance(query_result, pd.DataFrame):
        if query_result.empty:
            result_text = "No data found."
        else:
            result_text = f"Query returned {len(query_result)} rows:\n"
            result_text += query_result.to_string(index=False, max_rows=None)
    else:
        result_text = str(query_result)
    
    interpretation_prompt = f"""
        Original Question: {original_question}
        
        Query Results:
        {result_text}
        
        Please analyze these results and provide a clear, meaningful answer to the original question. 
        Consider ALL the data provided, not just the first few rows. 
        If there are patterns, summaries, or insights that can be drawn from the complete dataset, please include them.
        Provide specific numbers, names, or values from the results when relevant.
        Format your response in a conversational way that directly answers the user's question.
    """
    
    try:
        response = llm.invoke(interpretation_prompt)
        return response.content if hasattr(response, 'content') else str(response)
    except Exception as e:
        return f"Query executed successfully but interpretation failed: {str(e)}\n\nRaw results:\n{result_text}"

class DatabaseTool:
    """Custom database tool for the combined agent"""
    def __init__(self, engine, llm):
        self.engine = engine
        self.llm = llm
        self.db_uri = f"mysql+pymysql://{DB_CONFIG['user']}:{DB_CONFIG['password']}@{DB_CONFIG['host']}:{DB_CONFIG['port']}/{DB_CONFIG['database']}"
        self.sql_db = SQLDatabase.from_uri(
            self.db_uri, 
            sample_rows_in_table_info=5,
            max_string_length=10000
        )
        self.toolkit = SQLDatabaseToolkit(db=self.sql_db, llm=llm)
        
        # Create database sub-agent
        template="""
            You are a SQL expert. Answer questions by writing and executing SQL queries using the tools below:
            {tools}

            INSTRUCTIONS:
            1. Write queries that return complete answers.
            2. Use COUNT(*) or aggregates when needed.
            3. Don't use LIMIT unless asked.
            4. Analyze the full result before answering.
            5. Give specific numbers, patterns, and insights.

            Follow this format:
            Question: the user's question
            Thought: your reasoning
            Action: one of [{tool_names}]
            Action Input: input to the action
            Observation: result of the action
            ... (repeat Thought/Action/Input/Observation as needed)
            Thought: I now know the final answer
            Final Answer: your complete answer with data insights

            Begin!
            Question: {input}
            Thought: {agent_scratchpad}
        """
        sql_prompt = PromptTemplate(
            template=template,
            input_variables=["input", "agent_scratchpad", "tools", "tool_names"]
        )
        
        self.db_agent = create_react_agent(llm, self.toolkit.get_tools(), sql_prompt)
        self.db_agent_executor = AgentExecutor(
            agent=self.db_agent, 
            tools=self.toolkit.get_tools(), 
            handle_parsing_errors=True,
            max_iterations=10,
            max_execution_time=120
        )
    
    def query_database(self, question):
        """Execute database query using the sub-agent"""
        try:
            # Check if it's a direct SQL command
            if self.is_direct_sql(question):
                result = execute_direct_query(self.engine, question)
                if isinstance(result, pd.DataFrame):
                    return interpret_query_results(self.llm, question, result)
                else:
                    return str(result)
            else:
                # Use the database agent for natural language queries
                result = self.db_agent_executor.invoke({"input": question})
                return result["output"] if isinstance(result, dict) and "output" in result else str(result)
        except Exception as e:
            return f"Database query failed: {str(e)}"
    
    def is_direct_sql(self, question):
        """Check if the question is a direct SQL command"""
        sql_commands = ["select", "insert", "update", "delete", "show", "describe", "create", "drop", "alter"]
        return any(question.lower().strip().startswith(cmd) for cmd in sql_commands)

def create_combined_agent_executor(vectorstore):
    """Create a combined agent that can use both RAG and database tools."""
    llm=create_ollama_llm()
    engine=setup_sqlalchemy_engine()
    db_tool_instance=DatabaseTool(engine,llm)
    retriever=vectorstore.as_retriever(k=4)

    def extract_section_from_content(lines):
        """Extract likely section header from content lines."""
        for line in lines[:3]:
            line=line.strip()
            if 5<len(line)<80 and (
                line.isupper() or 
                re.match(r'^\d+\.',line) or 
                re.match(r'^(Chapter|Section)\b',line,re.IGNORECASE) or
                re.match(r'^[A-Z][^.!?]*:$',line) or 
                line.startswith('#')
            ):
                return line.replace('#','').strip()
        return "Document Content"

    def rag_search(question):
        """Enhanced RAG search with section & source tracking."""
        try:
            docs=retriever.invoke(question)
            if not docs:
                return "No relevant documents found."
            context_parts=[]
            source_info=set()
            for doc in docs:
                source=doc.metadata.get("source","Unknown Source")
                page=doc.metadata.get("page",doc.metadata.get("page_number","N/A"))
                section=extract_section_from_content(doc.page_content.split('\n'))
                context_parts.append(doc.page_content)
                source_info.add(f"{source} - {section} (Page {page})")
            context="\n\n".join(context_parts)
            sources="Sources: "+"; ".join(source_info)
            rag_prompt=(
                f"Answer the question using only the context below.\n\n"
                f"Question: {question}\n\n"
                f"Context:\n{context}\n\n"
                f"After your answer, include the following source information:\n"
                f"{sources}"
            )
            response=llm.invoke([{"role":"system","content":rag_prompt}])
            answer=getattr(response,"content",str(response))
            if "Sources:" not in answer:
                answer+=f"\n\n{sources}"
            return answer
        except Exception as e:
            return f"RAG search failed: {e}"

    # Define tools
    database_tool=Tool(
        name="DatabaseQuery",
        func=db_tool_instance.query_database,
        description=(
            "Use this tool to query the database. Best for structured data tasks such as:\n"
            "- Products, prices, invoices, customers, contacts\n"
            "- Keywords: count, sum, average, records, data, table, database"
        )
    )

    rag_tool=Tool(
        name="DocumentSearch",
        func=rag_search,
        description=(
            "Use this tool to search uploaded documents for:\n"
            "- Concepts, procedures, policies, explanations\n"
            "- Knowledge-base or text-based retrieval"
        )
    )

    template=("""You are an intelligent assistant with access to both a database and a document collection.
        TOOLS:\n{tools}\n
        ROUTING GUIDELINES:\n
        1. Use DatabaseQuery for structured or numerical data (e.g., products, prices, customers, contacts, invoices, SQL-like queries).\n
        2. Use DocumentSearch for conceptual or descriptive information (e.g., explanations, procedures, policies, meanings).\n
        3. Use both tools when:
        - The question needs both structured data and conceptual info
        - One tool's result leads to querying the other
        - Additional context is needed for a complete answer
        - One tool's result means 'I don't know' or 'I cannot answer' or 'insufficient context'.\n
        4. Always include the source of answers retrieved (section, page, or table name).\n
        5. Stop once you have a complete, well-supported final answer.\n\n
        FORMAT:\n
        Question: {input}\nThought: {agent_scratchpad}\n
        Action: one of [{tool_names}]\nAction Input: input to the action\nObservation: result\n
        ... (Repeat up to 5 times)\n
        Thought: I now know the final answer\nFinal Answer: combine insights from all tools\n
        Begin!\nQuestion: {input}\nThought: {agent_scratchpad}
        """
    )

    prompt=PromptTemplate(
        template=template,
        input_variables=["input","agent_scratchpad","tools","tool_names"]
    )

    agent=create_react_agent(
        llm=llm,
        tools=[database_tool,rag_tool],
        prompt=prompt
    )

    agent_executor=AgentExecutor(
        agent=agent,
        tools=[database_tool,rag_tool],
        handle_parsing_errors=True,
        max_iterations=10,
        max_execution_time=300,
        verbose=True
    )

    return agent_executor

# RFP Generation Functions
def extract_all_document_content(directory_path):
    """Extract all content from documents without chunking for RFP generation"""
    directory_path = Path(directory_path)
    file_extensions = ["*.pdf", "*.docx", "*.csv", "*.txt",
                       "*.jpg", "*.jpeg", "*.png", "*.gif", "*.bmp", "*.tiff",
                       "*.mp3", "*.wav", "*.flac", "*.m4a", "*.ogg",
                       "*.mp4", "*.mov", "*.avi"]
    all_files = []
    for extension in file_extensions:
        all_files.extend(directory_path.glob(extension))
    documents_content = {}
    for file in all_files:
        try:
            file_extension = file.suffix.lower()
            if file_extension == ".pdf":
                content = extract_pdf_text(str(file))
            elif file_extension == ".docx":
                content = extract_docx_text(str(file))
            elif file_extension == ".csv":
                content = extract_csv_text(str(file))
            elif file_extension == ".txt":
                content = extract_txt_text(str(file))
            elif file_extension in [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff"]:
                from file_handler.image_handler import extract_image_description, extract_image_text_easyocr
                content = extract_image_text_easyocr(str(file))
                if not content or content.strip() == "":
                    content = extract_image_description(str(file))
            elif file_extension in [".mp3", ".wav", ".flac", ".m4a", ".ogg"]:
                from file_handler.audio_handler import extract_audio_text
                content = extract_audio_text(str(file))
            elif file_extension in [".mp4", ".mov", ".avi"]:
                from file_handler.video_handler import extract_video_content
                content = extract_video_content(str(file))
            else:
                continue   
            if content and content.strip():
                documents_content[file.name] = content
        except Exception as e:
            print(f"Error extracting content from {file.name}: {str(e)}")
    return documents_content

def get_database_summary(engine):
    """Get a comprehensive summary of database content for RFP"""
    try:
        # Get table information
        tables_query = "SHOW TABLES"
        tables_df = execute_direct_query(engine, tables_query)
        database_summary = {"tables": {}}
        for _, row in tables_df.iterrows():
            table_name = row.iloc[0]  # First column contains table name
            # Get table structure
            desc_query = f"DESCRIBE {table_name}"
            structure_df = execute_direct_query(engine, desc_query)
            # Get sample data (first 5 rows)
            sample_query = f"SELECT * FROM {table_name} LIMIT 5"
            sample_df = execute_direct_query(engine, sample_query)
            # Get row count
            count_query = f"SELECT COUNT(*) as count FROM {table_name}"
            count_df = execute_direct_query(engine, count_query)
            database_summary["tables"][table_name] = {
                "structure": structure_df.to_string(index=False) if not structure_df.empty else "No structure info",
                "sample_data": sample_df.to_string(index=False) if not sample_df.empty else "No sample data",
                "total_rows": count_df.iloc[0]['count'] if not count_df.empty else 0
            }
        return database_summary
    except Exception as e:
        return {"error": f"Failed to get database summary: {str(e)}"}

def generate_rfp(documents_content, database_summary, project_type="general", db_use="no", user_prompt=""):
    """Generate RFP using documents and database information OR user prompt"""
    llm = create_ollama_llm()
    
    # If user provided a prompt, use that instead of documents
    if user_prompt and user_prompt.strip():
        rfp_prompt = f"""
        You are an expert RFP (Request for Proposal) writer. Generate a comprehensive RFP document based on the user's description.

        USER DESCRIPTION:
        {user_prompt}

        PROJECT TYPE: {project_type}

        Generate a professional RFP document with the following sections:

        1. PROJECT OVERVIEW
        - Brief description of the project
        - Background information
        
        2. SCOPE OF WORK
        - Detailed requirements
        - Technical specifications
        
        3. DELIVERABLES
        - Expected outputs and deliverables
        - Timeline expectations
        
        4. TECHNICAL REQUIREMENTS
        - System requirements
        - Technology stack preferences
        
        5. EVALUATION CRITERIA
        - How proposals will be evaluated
        - Key success factors
        
        6. SUBMISSION REQUIREMENTS
        - Required proposal format
        - Submission deadline and process
        
        7. TERMS AND CONDITIONS
        - Basic contractual terms
        - Payment structure

        Make the RFP specific and detailed based on the user's description.
        """
    else:
        # Original document-based generation
        # Prepare document context
        doc_context = "\n\n".join([f"Document: {filename}\nContent: {content}" for filename, content in documents_content.items()])
        
        # Prepare database context
        db_context = ""
        if "tables" in database_summary:
            for table_name, table_info in database_summary["tables"].items():
                db_context += f"\nTable: {table_name}\n"
                db_context += f"Structure: {table_info['structure']}\n"
                db_context += f"Sample Data: {table_info['sample_data']}\n"
                db_context += f"Total Records: {table_info['total_rows']}\n"
        elif "error" in database_summary:
            db_context = database_summary["error"]
        
        # RFP generation prompt with DB
        rfp_prompt_with_db = f"""
        You are an expert RFP (Request for Proposal) writer. Generate a comprehensive RFP document based on the provided documents and database information.

        DOCUMENT CONTENT:
        {doc_context}

        DATABASE INFORMATION:
        {db_context}

        PROJECT TYPE: {project_type}

        Generate a professional RFP document with the following sections:

        1. PROJECT OVERVIEW
        - Brief description of the project
        - Background information from documents
        
        2. SCOPE OF WORK
        - Detailed requirements based on document analysis
        - Technical specifications
        - Data requirements (based on database structure)
        
        3. DELIVERABLES
        - Expected outputs and deliverables
        - Timeline expectations
        
        4. TECHNICAL REQUIREMENTS
        - System requirements
        - Database integration needs
        - Technology stack preferences
        
        5. EVALUATION CRITERIA
        - How proposals will be evaluated
        - Key success factors
        
        6. SUBMISSION REQUIREMENTS
        - Required proposal format
        - Submission deadline and process
        
        7. TERMS AND CONDITIONS
        - Basic contractual terms
        - Payment structure

        Make the RFP specific to the content provided in documents and database structure. Use concrete examples and requirements derived from the available information.
        """
        
        # RFP generation prompt without DB
        rfp_prompt_without_db = f"""
        You are an expert RFP (Request for Proposal) writer. Generate a comprehensive RFP document based on the provided documents.

        DOCUMENT CONTENT:
        {doc_context}

        PROJECT TYPE: {project_type}

        Generate a professional RFP document with the following sections:

        1. PROJECT OVERVIEW
        - Brief description of the project
        - Background information from documents
        
        2. SCOPE OF WORK
        - Detailed requirements based on document analysis
        - Technical specifications
        
        3. DELIVERABLES
        - Expected outputs and deliverables
        - Timeline expectations
        
        4. TECHNICAL REQUIREMENTS
        - System requirements
        - Technology stack preferences
        
        5. EVALUATION CRITERIA
        - How proposals will be evaluated
        - Key success factors
        
        6. SUBMISSION REQUIREMENTS
        - Required proposal format
        - Submission deadline and process
        
        7. TERMS AND CONDITIONS
        - Basic contractual terms
        - Payment structure

        Make the RFP specific to the content provided in documents. Use concrete examples and requirements derived from the available information.
        """
        
        rfp_prompt = rfp_prompt_with_db if db_use == "yes" else rfp_prompt_without_db
    
    try:
        response = llm.invoke(rfp_prompt)
        rfp_content = response.content if hasattr(response, 'content') else str(response)
        return rfp_content
    except Exception as e:
        return f"Error generating RFP: {str(e)}"

def save_rfp_to_file(rfp_content, filename="generated_rfp.txt"):
    """Save RFP content to a file"""
    try:
        output_path = Path("generated_rfps")
        output_path.mkdir(exist_ok=True)
        file_path = output_path / filename
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(rfp_content)
        return str(file_path)
    except Exception as e:
        return f"Error saving RFP: {str(e)}"

def remove_think_tags(text):
    cleaned_text = re.sub(r'<(think|thinking|reflection)>.*?</\1>', '', text, flags=re.DOTALL)
    return cleaned_text.strip()

def clean_markdown(text):
    """Clean markdown text so it shows as regular text"""
    # Remove bold and italics (**text**, *text*)
    text=re.sub(r'\*\*(.*?)\*\*',r'\1',text)
    text=re.sub(r'\*(.*?)\*',r'\1',text)
    # Remove inline code (`code`)
    text=re.sub(r'`(.*?)`',r'\1',text)
    # Remove links but keep text [text](url)
    text=re.sub(r'\[(.*?)\]\(.*?\)',r'\1',text)
    # Remove images ![alt](url)
    text=re.sub(r'!\[(.*?)\]\(.*?\)',r'\1',text)
    # Remove extra # used for headings
    text=re.sub(r'^#+\s*','',text)
    return text.strip()

def convert_to_pdf(content):
    """Convert markdown content to PDF using reportlab after removing markdown syntax"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        import io
        # Function to clean markdown
        buffer=io.BytesIO()
        doc=SimpleDocTemplate(buffer,pagesize=letter)
        styles=getSampleStyleSheet()
        story=[]
        lines=content.split('\n')
        for line in lines:
            clean_line=clean_markdown(line)
            if clean_line:
                if line.strip().startswith('#'):
                    story.append(Paragraph(clean_line,styles['Heading1']))
                else:
                    story.append(Paragraph(clean_line,styles['Normal']))
            story.append(Spacer(1,0.1*inch))
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        return None

def convert_to_docx(content):
    """Convert markdown content to DOCX using python-docx"""
    try:
        doc=docx_Doc()
        lines=content.split('\n')
        for line in lines:
            clean_line=clean_markdown(line)
            if clean_line:
                if line.strip().startswith('#'):
                    doc.add_heading(clean_line,level=1)
                else:
                    doc.add_paragraph(clean_line)
        buffer=io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        return None

def cleanup_temp_files():
    """Clean up temporary files and directories"""
    try:
        if os.path.exists(FOLDER):
            shutil.rmtree(FOLDER)
        if os.path.exists(CHROMA_DIR):
            shutil.rmtree(CHROMA_DIR)
        Path(FOLDER).mkdir(parents=True, exist_ok=True)
        Path(CHROMA_DIR).mkdir(parents=True, exist_ok=True)
    except Exception as e:
        print(f"Error cleaning up temp files: {str(e)}")

def run_rfp_generation():
    """Run RFP generation workflow"""
    print("\n" + "="*50)
    print("RFP GENERATION MODE")
    print("="*50)
    try:
        # Extract document content
        print("Extracting content from documents...")
        documents_content = extract_all_document_content(FOLDER)
        if not documents_content:
            print("No documents found or no content extracted.")
            return
        print(f"Extracted content from {len(documents_content)} documents")
        # Get database summary
        print("Analyzing database...")
        engine = setup_sqlalchemy_engine()
        database_summary = get_database_summary(engine)
        # Get project type from user
        project_type = input("Enter project type (or press Enter for 'general'): ").strip()
        if not project_type:
            project_type = "general"
        db_use=input("Use connected DB for RFP? Write 'yes' (or press Enter for 'no'): ").strip()
        if not project_type:
            db_use = "no"
        # Generate RFP
        print("Generating RFP document...")
        rfp_content = generate_rfp(documents_content, database_summary, project_type, db_use)
        # Display RFP
        print("\n" + "="*50)
        print("GENERATED RFP")
        print("="*50)
        print(rfp_content)
        
        # Ask if user wants to save
        save_option = input("\nDo you want to save this RFP to a file? (y/n): ").strip().lower()
        if save_option == 'y':
            filename = input("Enter filename (or press Enter for default): ").strip()
            if not filename:
                filename = f"rfp_{project_type}_{int(time.time())}.txt"
            elif not filename.endswith('.txt'):
                filename += '.txt'
            rfp_content=remove_think_tags(rfp_content)
            file_path = save_rfp_to_file(rfp_content, filename)
            print(f"RFP saved to: {file_path}")
    except Exception as e:
        print(f"Error in RFP generation: {str(e)}")

def main():
    """Main function with mode selection"""
    print("Setting up system...")
    try:
        vectorstore = setup_rag_vectorstore(FOLDER)
        print("System ready!")
        while True:
            print("\n" + "="*50)
            print("SELECT MODE:")
            print("1. Chatbot (Query documents and database)")
            print("2. Generate RFP")
            print("3. Quit")
            print("="*50)
            choice = input("Enter your choice (1/2/3): ").strip()
            if choice == "1":
                # Chatbot mode
                agent_executor = create_combined_agent_executor(vectorstore)
                print("\nChatbot mode activated!")
                print("Commands: Ask any question, 'back' to return to menu")
                while True:
                    question = input("\nAsk a question (or 'back' for main menu): ").strip()
                    if question.lower() == "back":
                        break
                    elif question:
                        try:
                            result = agent_executor.invoke({"input": question})
                            answer = result["output"]
                            print(f"\nAnswer: {answer}")
                        except Exception as e:
                            print(f"Error processing question: {str(e)}")
                            print("Please try rephrasing your question.")
            elif choice == "2":
                # RFP generation mode
                run_rfp_generation()
            elif choice == "3":
                print("Goodbye!")
                break
            else:
                print("Invalid choice. Please select 1, 2, or 3.")
    except Exception as e:
        print(f"Error setting up system: {str(e)}")
        print("Make sure Ollama is running and the specified model is available.")

if __name__ == "__main__":
    main()
