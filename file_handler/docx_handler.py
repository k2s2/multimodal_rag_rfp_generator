from docx import Document
import zipfile
import io
import os
import numpy as np
import easyocr
from PIL import Image
from file_handler.image_handler import extract_image_description


# New optimised code
def extract_docx_text(docx_path):
    """Extract text from DOCX using EasyOCR for images, with fallback to image description"""
    reader = easyocr.Reader(['en'])
    doc = Document(docx_path)
    all_text = []
    # Extract regular text
    full_text = [p.text for p in doc.paragraphs if p.text.strip()]
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(' | '.join(row_text))
    if full_text:
        all_text.append("Document Text:\n" + '\n'.join(full_text))
    # Extract and OCR images
    try:
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            image_extensions = ('.png', '.jpg', '.jpeg', '.gif', '.bmp')
            media_files = [f for f in docx_zip.namelist() 
                          if f.startswith('word/media/') and f.lower().endswith(image_extensions)]
            for img_index, media_file in enumerate(media_files):
                try:
                    img_data = docx_zip.read(media_file)
                    pil_image = Image.open(io.BytesIO(img_data))
                    if pil_image.mode in ('RGBA', 'P'):
                        pil_image = pil_image.convert('RGB')
                    img_array = np.array(pil_image)
                    results = reader.readtext(img_array)
                    ocr_text = '\n'.join([result[1] for result in results])
                    if ocr_text.strip():
                        all_text.append(f"Image {img_index + 1} (OCR):\n{ocr_text}")
                    else:
                        try:
                            temp_path = f"temp_docx_img_{img_index}.png"
                            pil_image.save(temp_path)
                            description = extract_image_description(temp_path)
                            os.remove(temp_path)
                            if description and description.strip():
                                all_text.append(f"Image {img_index + 1} (Description):\n{description}")
                        except Exception:
                            pass
                except Exception:
                    continue
    except Exception:
        pass
    
    return '\n\n'.join(all_text)

# def extract_docx_text(docx_path):
#     """Extract text from DOCX using EasyOCR for images, with fallback to image description"""
#     reader = easyocr.Reader(['en'])
#     doc = Document(docx_path)
#     all_text = []
#     # Extract regular text
#     full_text = []
#     for paragraph in doc.paragraphs:
#         if paragraph.text.strip():
#             full_text.append(paragraph.text)
#     # Extract text from tables
#     for table in doc.tables:
#         for row in table.rows:
#             row_text = []
#             for cell in row.cells:
#                 if cell.text.strip():
#                     row_text.append(cell.text.strip())
#             if row_text:
#                 full_text.append(' | '.join(row_text))
#     if full_text:
#         all_text.append("Document Text:\n" + '\n'.join(full_text))
#     # Extract and OCR images
#     try:
#         with zipfile.ZipFile(docx_path, 'r') as docx_zip:
#             media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
#             for img_index, media_file in enumerate(media_files):
#                 try:
#                     if not any(media_file.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']):
#                         continue
#                     img_data = docx_zip.read(media_file)
#                     pil_image = Image.open(io.BytesIO(img_data))
#                     if pil_image.mode in ('RGBA', 'P'):
#                         pil_image = pil_image.convert('RGB')
#                     img_array = np.array(pil_image)
#                     # Use EasyOCR first
#                     results = reader.readtext(img_array)
#                     ocr_text = '\n'.join([result[1] for result in results]) if results else ""
#                     if ocr_text and ocr_text.strip():
#                         # OCR found text
#                         all_text.append(f"Image {img_index + 1} (OCR):\n{ocr_text}")
#                     else:
#                         # No text found via OCR, try image description
#                         try:
#                             # Save temporary image for description function
#                             temp_img_path = f"temp_docx_img_{img_index}.png"
#                             pil_image.save(temp_img_path)
#                             # Get visual description
#                             description = extract_image_description(temp_img_path)
#                             if description and description.strip():
#                                 all_text.append(f"Image {img_index + 1} (Description):\n{description}")
#                             # Clean up temp file
#                             if os.path.exists(temp_img_path):
#                                 os.remove(temp_img_path)
#                         except Exception as desc_error:
#                             print(f"Error getting image description: {desc_error}")
#                 except Exception as e:
#                     print(f"Error processing image: {e}")
#     except Exception as e:
#         print(f"Error accessing DOCX file: {e}")
#     return '\n\n'.join(all_text)